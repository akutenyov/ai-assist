"""Console chat client powered by Groq Compound with live web search."""

from __future__ import annotations

import json
import logging
import sys
from datetime import date, datetime
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Final


def remove_external_codex_dependencies() -> None:
    """Avoid packages injected by another Codex/PyCharm project into sys.path.

    This application must load its dependencies from the selected virtual
    environment. A foreign ``_codex_pydeps`` folder can otherwise override
    pydantic and cause ``pydantic_core`` import errors before Groq starts.
    """
    sys.path[:] = [
        entry
        for entry in sys.path
        if Path(entry).name.lower() != "_codex_pydeps"
    ]


remove_external_codex_dependencies()

from groq import APIConnectionError, APIStatusError, Groq, RateLimitError
from rich import box
from rich.console import Console
from rich.markdown import Markdown
from rich.panel import Panel
from rich.table import Table
from rich.text import Text

from config import GROQ_API_KEY, GROQ_MODEL, SYSTEM_PROMPT


MAX_CONTEXT_MESSAGES = 12
MAX_MESSAGE_CHARS = 6_000
MAX_SEARCH_SOURCES = 5
CHAT_HISTORY_DIRECTORY = "chats"
MAX_FILE_BYTES: Final = 5 * 1024 * 1024
MAX_FILE_TEXT_CHARS: Final = 4_500
MAX_FILE_PROMPT_CHARS: Final = 800
MAX_SPREADSHEET_ROWS: Final = 200
SUPPORTED_FILE_TYPES: Final = {
    ".txt": "текстовий файл",
    ".md": "Markdown",
    ".py": "Python-код",
    ".json": "JSON",
    ".csv": "CSV-таблиця",
    ".log": "журнал",
    ".yaml": "YAML",
    ".yml": "YAML",
    ".xml": "XML",
    ".html": "HTML",
    ".pdf": "PDF із текстовим шаром",
    ".docx": "Word DOCX",
    ".xlsx": "Excel XLSX",
}
WEEKDAYS_UA = (
    "понеділок",
    "вівторок",
    "середа",
    "четвер",
    "п’ятниця",
    "субота",
    "неділя",
)


def configure_logging() -> logging.Logger:
    """Create a rotating local log without storing prompts or API secrets."""
    logger = logging.getLogger("ai_assist")

    if logger.handlers:
        return logger

    log_directory = Path(__file__).resolve().parent / "logs"
    log_directory.mkdir(exist_ok=True)

    handler = RotatingFileHandler(
        log_directory / "ai_assist.log",
        maxBytes=1_000_000,
        backupCount=3,
        encoding="utf-8",
    )
    handler.setFormatter(
        logging.Formatter(
            "%(asctime)s | %(levelname)s | %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S",
        )
    )

    logger.setLevel(logging.INFO)
    logger.addHandler(handler)
    logger.propagate = False

    return logger


LOGGER = configure_logging()

# Some legacy Windows consoles still use CP1251. Replacing unsupported glyphs
# avoids a fatal UnicodeEncodeError while preserving all Ukrainian text.
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(errors="replace")

CONSOLE = Console(emoji=False)


def validate_configuration() -> None:
    """Stop before an API call when the required credential is missing."""
    if not GROQ_API_KEY:
        raise RuntimeError(
            "GROQ_API_KEY не заданий. Додайте ключ до файлу .env."
        )


def build_system_prompt() -> str:
    """Add the current local date so ambiguous dates are interpreted correctly."""
    today = date.today()
    current_date = today.strftime("%d.%m.%Y")
    weekday = WEEKDAYS_UA[today.weekday()]

    return (
        f"{SYSTEM_PROMPT}\n\n"
        f"Поточна дата: {current_date}, {weekday}. "
        "Інтерпретуй неповні дати (наприклад, '16-08') відносно цієї "
        "дати. Для прогнозів і новин перевіряй актуальні дані веб-пошуком."
    )


def truncate_text(text: str, limit: int = MAX_MESSAGE_CHARS) -> str:
    """Keep a message within the API context budget without losing both ends."""
    if len(text) <= limit:
        return text

    marker = "\n...[повідомлення скорочено]...\n"
    remaining = limit - len(marker)
    start_length = remaining // 2
    end_length = remaining - start_length

    return f"{text[:start_length]}{marker}{text[-end_length:]}"


def build_request_messages(
    messages: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Build a bounded API payload with a fresh date-aware system prompt."""
    dialogue = messages[1:][-MAX_CONTEXT_MESSAGES:]

    return [
        {"role": "system", "content": truncate_text(build_system_prompt())},
        *[
            {
                "role": message["role"],
                "content": truncate_text(message["content"]),
            }
            for message in dialogue
        ],
    ]


def new_conversation() -> list[dict[str, str]]:
    """Create a fresh dialogue with a date-aware system message."""
    return [{"role": "system", "content": build_system_prompt()}]


def create_chat_history_path() -> Path:
    """Create a unique JSON path for the current chat session."""
    directory = Path(__file__).resolve().parent / CHAT_HISTORY_DIRECTORY
    directory.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S_%f")

    return directory / f"chat_{timestamp}.json"


def save_chat_history(
    history_path: Path,
    messages: list[dict[str, str]],
) -> None:
    """Atomically save the local chat without storing any API credentials."""
    payload = {
        "format_version": 1,
        "saved_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "model": GROQ_MODEL,
        "messages": messages,
    }
    temporary_path = history_path.with_suffix(".tmp")
    temporary_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    temporary_path.replace(history_path)


def print_saved_chats() -> None:
    """List saved chat files without reading their private contents."""
    directory = Path(__file__).resolve().parent / CHAT_HISTORY_DIRECTORY
    chat_files = sorted(
        directory.glob("chat_*.json") if directory.exists() else [],
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )

    if not chat_files:
        CONSOLE.print("[dim]Збережених чатів ще немає.[/]")
        return

    table = Table(
        title="[bold cyan]Збережені чати[/]",
        box=box.SIMPLE_HEAVY,
        header_style="bold cyan",
    )
    table.add_column("Файл", style="white")
    table.add_column("Останнє збереження", style="dim")

    for path in chat_files[:10]:
        saved_at = datetime.fromtimestamp(path.stat().st_mtime).strftime(
            "%d.%m.%Y %H:%M"
        )
        table.add_row(path.name, saved_at)

    CONSOLE.print(table)


def print_file_capabilities() -> None:
    """Explain supported local files before any content is sent to Groq."""
    table = Table(
        title="[bold cyan]Обробка файлів[/]",
        box=box.SIMPLE_HEAVY,
        header_style="bold cyan",
    )
    table.add_column("Формат", style="bold yellow")
    table.add_column("Що буде опрацьовано")
    table.add_row("TXT, MD, PY, JSON, CSV, LOG, YAML, XML, HTML", "Текст і код")
    table.add_row("PDF", "Текстовий шар; скановані PDF без OCR не підтримуються")
    table.add_row("DOCX", "Текст документа")
    table.add_row("XLSX", f"Значення комірок, до {MAX_SPREADSHEET_ROWS} рядків")
    CONSOLE.print(table)
    CONSOLE.print(
        Panel(
            Text.from_markup(
                f"[bold yellow]Обмеження:[/] файл до {MAX_FILE_BYTES // 1024 // 1024} МБ; "
                f"до {MAX_FILE_TEXT_CHARS:,} символів тексту передається моделі.\n"
                "Фото, аудіо й відео не підтримуються поточною текстовою моделлю "
                f"[cyan]{GROQ_MODEL}[/].\n"
                "Вміст вибраного файла буде надіслано до Groq разом із вашим наступним "
                "запитом і збережено у локальній історії чату. Не додавайте секретні дані."
            ),
            title="[bold yellow]Перед додаванням файла[/]",
            border_style="yellow",
            padding=(1, 2),
        )
    )


def decode_text_file(file_path: Path) -> str:
    """Read common text files while handling typical Ukrainian encodings."""
    raw_content = file_path.read_bytes()

    if b"\x00" in raw_content:
        raise ValueError("Файл схожий на двійковий, а не текстовий.")

    for encoding in ("utf-8-sig", "utf-8", "cp1251", "utf-16"):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("Не вдалося визначити кодування текстового файла.")


def extract_pdf_text(file_path: Path) -> str:
    """Extract the text layer from a PDF without uploading the original file."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]

    return "\n\n".join(pages).strip()


def extract_docx_text(file_path: Path) -> str:
    """Extract paragraphs and table cells from a Word document."""
    from docx import Document

    document = Document(file_path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cell for cell in cells if cell))

    return "\n".join(part for part in parts if part).strip()


def extract_xlsx_text(file_path: Path) -> str:
    """Extract a bounded number of cell values from every non-empty worksheet."""
    from openpyxl import load_workbook

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    parts: list[str] = []
    rows_read = 0

    try:
        for worksheet in workbook.worksheets:
            if rows_read >= MAX_SPREADSHEET_ROWS:
                break

            parts.append(f"Аркуш: {worksheet.title}")

            for row in worksheet.iter_rows(values_only=True):
                if rows_read >= MAX_SPREADSHEET_ROWS:
                    break

                values = [str(value).strip() if value is not None else "" for value in row]

                if any(values):
                    parts.append(" | ".join(values))
                    rows_read += 1
    finally:
        workbook.close()

    return "\n".join(parts).strip()


def extract_file_text(file_path: Path) -> str:
    """Select a local extractor according to a pre-approved file extension."""
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    if extension == ".docx":
        return extract_docx_text(file_path)

    if extension == ".xlsx":
        return extract_xlsx_text(file_path)

    return decode_text_file(file_path)


def prepare_file_attachment(file_path: Path) -> str:
    """Validate a local file and make a bounded text attachment for the model."""
    if not file_path.exists() or not file_path.is_file():
        raise ValueError("Файл не знайдено або це не звичайний файл.")

    extension = file_path.suffix.lower()

    if extension not in SUPPORTED_FILE_TYPES:
        supported = ", ".join(SUPPORTED_FILE_TYPES)
        raise ValueError(f"Непідтримуваний формат. Доступні: {supported}.")

    file_size = file_path.stat().st_size

    if file_size > MAX_FILE_BYTES:
        raise ValueError(
            f"Файл має {file_size / 1024 / 1024:.1f} МБ. Максимум — "
            f"{MAX_FILE_BYTES // 1024 // 1024} МБ."
        )

    extracted_text = extract_file_text(file_path)
    extracted_text = "\n".join(line.rstrip() for line in extracted_text.splitlines())

    if not extracted_text.strip():
        raise ValueError(
            "Не вдалося отримати текст. Для PDF це може бути скан без "
            "текстового шару."
        )

    attachment_text = truncate_text(extracted_text.strip(), MAX_FILE_TEXT_CHARS)
    return (
        f"[Прикріплений файл: {file_path.name}; формат: {extension}; "
        f"розмір: {file_size / 1024:.1f} КБ]\n"
        "Вміст файла:\n"
        f"{attachment_text}"
    )


def build_file_request(prompt: str, attachment: str) -> str:
    """Keep the user task and an already-bounded file within one API message."""
    task = truncate_text(prompt, MAX_FILE_PROMPT_CHARS)
    return (
        f"Завдання користувача щодо прикріпленого файла:\n{task}\n\n"
        f"{attachment}\n\n"
        "Відповідай за змістом файла. Якщо даних недостатньо, прямо скажи це."
    )


def read_field(value: object, field: str) -> object | None:
    """Read an SDK object or dictionary field without relying on one shape."""
    if isinstance(value, dict):
        return value.get(field)

    return getattr(value, field, None)


def extract_search_sources(executed_tools: list[object]) -> list[tuple[str, str]]:
    """Extract up to five unique titles and URLs from Groq web-search results."""
    sources: list[tuple[str, str]] = []
    seen_urls: set[str] = set()

    for tool in executed_tools:
        if read_field(tool, "type") not in {"search", "web_search"}:
            continue

        results = read_field(tool, "search_results") or []

        if isinstance(results, dict):
            results = results.get("results", [])

        for result in results:
            url = str(read_field(result, "url") or "").strip()

            if not url or url in seen_urls:
                continue

            title = str(read_field(result, "title") or url)
            title = " ".join(title.split())[:160]
            sources.append((title, url))
            seen_urls.add(url)

            if len(sources) == MAX_SEARCH_SOURCES:
                return sources

    return sources


def get_header(headers: object, name: str) -> str | None:
    """Read an HTTP header from the SDK response, ignoring its letter case."""
    getter = getattr(headers, "get", None)

    if callable(getter):
        value = getter(name)

        if value is not None:
            return str(value)

    items = getattr(headers, "items", None)

    if callable(items):
        for header_name, value in items():
            if str(header_name).lower() == name.lower():
                return str(value)

    return None


def extract_rate_limits(headers: object) -> list[str]:
    """Format the remaining Groq limits included in HTTP response headers.

    Groq can change the active rate-limit window depending on the account and
    model. Therefore this function shows the exact limit reported by the API
    instead of assuming that it is a daily quota.
    """
    limits: list[str] = []

    for label, suffix in (("Запити", "requests"), ("Токени", "tokens")):
        remaining = get_header(headers, f"x-ratelimit-remaining-{suffix}")
        limit = get_header(headers, f"x-ratelimit-limit-{suffix}")
        reset_after = get_header(headers, f"x-ratelimit-reset-{suffix}")

        if remaining is None:
            continue

        line = f"{label}: залишилось {remaining}"

        if limit:
            line += f" з {limit}"

        if reset_after:
            line += f"; оновлення через {reset_after}"

        limits.append(line)

    return limits


def print_help() -> None:
    """Show the commands supported by the console chat."""
    table = Table(
        title="[bold cyan]Команди[/]",
        box=box.SIMPLE_HEAVY,
        header_style="bold cyan",
    )
    table.add_column("Команда", style="bold yellow")
    table.add_column("Дія")
    table.add_row("/help", "Показати цю довідку.")
    table.add_row("/clear", "Почати новий діалог без попереднього контексту.")
    table.add_row("/history", "Показати десять останніх збережених чатів.")
    table.add_row("/file", "Показати формати й обмеження завантаження.")
    table.add_row("/file <шлях>", "Додати текст із локального файла до наступного запиту.")
    table.add_row("/remove-file", "Прибрати підготовлений файл із наступного запиту.")
    table.add_row("/exit", "Завершити програму.")
    CONSOLE.print(table)
    CONSOLE.print(
        "[dim]Модель застосовує веб-пошук, коли потрібні актуальні дані.[/]"
    )


def print_banner(history_path: Path) -> None:
    """Show compact startup information in a styled panel."""
    content = Text()
    content.append("Модель: ", style="bold")
    content.append(GROQ_MODEL, style="cyan")
    content.append("\nВеб-пошук: ", style="bold")
    content.append("автоматично", style="green")
    content.append("\nІсторія: ", style="bold")
    content.append(f"{CHAT_HISTORY_DIRECTORY}\\{history_path.name}", style="dim")
    content.append("\n/help — список команд", style="dim")
    CONSOLE.print(
        Panel(
            content,
            title="[bold bright_cyan]AI Assist[/]",
            border_style="bright_cyan",
            padding=(1, 2),
        )
    )


def print_assistant_reply(reply: str) -> None:
    """Render the model reply as Markdown in a clearly separated panel."""
    CONSOLE.print(
        Panel(
            Markdown(reply),
            title="[bold green]Groq[/]",
            border_style="green",
            padding=(1, 2),
        )
    )


def print_sources(sources: list[tuple[str, str]]) -> None:
    """Render web-search sources in a compact clickable table."""
    table = Table(
        title="[bold cyan]Джерела веб-пошуку[/]",
        box=box.SIMPLE_HEAVY,
        header_style="bold cyan",
    )
    table.add_column("#", style="dim", width=3)
    table.add_column("Джерело", style="white")
    table.add_column("URL", style="cyan")

    for number, (title, url) in enumerate(sources, start=1):
        url_text = Text(url, style=f"link {url} underline cyan")
        table.add_row(str(number), Text(title), url_text)

    CONSOLE.print(table)


def print_rate_limits(rate_limits: list[str]) -> None:
    """Render the latest Groq quota values returned by the API."""
    if not rate_limits:
        CONSOLE.print("[dim]Квота Groq: API не передав ці дані.[/]")
        return

    table = Table.grid(padding=(0, 1))

    for rate_limit in rate_limits:
        table.add_row("[bold magenta]-[/]", Text(rate_limit))

    CONSOLE.print(
        Panel(
            table,
            title="[bold magenta]Залишок квоти Groq[/]",
            border_style="magenta",
            padding=(0, 1),
        )
    )


def print_error(message: str, title: str = "Помилка") -> None:
    """Show an error without interpreting API text as Rich markup."""
    CONSOLE.print(
        Panel(
            Text(message),
            title=f"[bold red]{title}[/]",
            border_style="red",
            padding=(0, 1),
        )
    )


def request_response(
    client: Groq,
    messages: list[dict[str, str]],
) -> tuple[str, bool, list[tuple[str, str]], list[str]]:
    """Send bounded history and return reply, sources, and live quota details.

    Compound automatically decides when live web search is necessary. The tool
    runs on Groq's servers; the program needs no browser.
    """
    raw_response = client.chat.completions.with_raw_response.create(
        model=GROQ_MODEL,
        messages=build_request_messages(messages),
        # Compound Mini automatically invokes its single built-in web-search
        # tool only when the question needs fresh information. Keeping this
        # request minimal avoids passing extra tool configuration to the API.
    )
    response = raw_response.parse()

    message = response.choices[0].message
    text = (message.content or "").strip()

    if not text:
        text = "Groq не повернув текстової відповіді."

    executed_tools = getattr(message, "executed_tools", None) or []
    web_search_used = any(
        read_field(tool, "type") in {"search", "web_search"}
        for tool in executed_tools
    )
    sources = extract_search_sources(executed_tools)
    rate_limits = extract_rate_limits(raw_response.headers)

    return text, web_search_used, sources, rate_limits


def run_chat() -> None:
    """Run the interactive terminal interface until the user exits."""
    validate_configuration()
    client = Groq(
        api_key=GROQ_API_KEY,
        # This stable Compound version uses basic web search. It is more
        # suitable for the free plan than the larger "latest" tool workflow.
        default_headers={"Groq-Model-Version": "2025-07-23"},
    )
    messages = new_conversation()
    history_path = create_chat_history_path()
    pending_attachment: str | None = None

    print_banner(history_path)
    LOGGER.info("Application started; model=%s", GROQ_MODEL)

    while True:
        try:
            prompt = CONSOLE.input("\n[bold cyan]Ви[/] [dim]>[/] ").strip()
        except (EOFError, KeyboardInterrupt):
            CONSOLE.print("\n[bold cyan]До побачення![/]")
            LOGGER.info("Application stopped by keyboard or input end")
            return

        if not prompt:
            continue

        command = prompt.lower()

        if command in {"/exit", "/quit", "/q"}:
            CONSOLE.print("[bold cyan]До побачення![/]")
            LOGGER.info("Application stopped by user command")
            return

        if command in {"/help", "help"}:
            print_help()
            continue

        if command in {"/clear", "clear"}:
            messages = new_conversation()
            history_path = create_chat_history_path()
            pending_attachment = None
            CONSOLE.print(
                "[bold green](OK) Почато новий діалог.[/] "
                f"[dim]Історія: {CHAT_HISTORY_DIRECTORY}\\{history_path.name}[/]"
            )
            LOGGER.info("Conversation context cleared; new chat history started")
            continue

        if command in {"/history", "history"}:
            print_saved_chats()
            continue

        if command in {"/file", "file"}:
            print_file_capabilities()
            CONSOLE.print("[dim]Приклад: /file D:\\Документи\\звіт.pdf[/]")
            continue

        if prompt.lower().startswith("/file "):
            print_file_capabilities()
            raw_path = prompt[6:].strip().strip('"').strip("'")

            if not raw_path:
                print_error("Вкажіть шлях після команди /file.", "Шлях не вказано")
                continue

            confirmation = CONSOLE.input(
                "[bold yellow]Прочитати цей файл і підготувати його для Groq?[/] "
                "\\[y/N] "
            ).strip().lower()

            if confirmation not in {"y", "yes", "т", "так"}:
                CONSOLE.print("[dim]Додавання файла скасовано.[/]")
                continue

            try:
                pending_attachment = prepare_file_attachment(Path(raw_path))
            except (OSError, ValueError) as error:
                print_error(str(error), "Файл не додано")
                LOGGER.warning("File attachment rejected: %s", error)
                continue
            except Exception as error:
                print_error(
                    f"Не вдалося обробити файл: {error}",
                    "Помилка обробки файла",
                )
                LOGGER.exception("Unable to prepare file attachment")
                continue

            CONSOLE.print(
                "[bold green](OK) Файл підготовлено.[/] Введіть наступний запит: "
                "наприклад, «Стисло підсумуй документ»."
            )
            LOGGER.info("File attachment prepared; suffix=%s", Path(raw_path).suffix.lower())
            continue

        if command in {"/remove-file", "remove-file"}:
            if pending_attachment is None:
                CONSOLE.print("[dim]Немає підготовленого файла.[/]")
            else:
                pending_attachment = None
                CONSOLE.print("[bold green](OK) Файл прибрано з наступного запиту.[/]")
                LOGGER.info("Pending file attachment removed")
            continue

        try:
            prompt_for_context = (
                build_file_request(prompt, pending_attachment)
                if pending_attachment
                else truncate_text(prompt)
            )

            if prompt_for_context != prompt:
                if pending_attachment and len(prompt) > MAX_FILE_PROMPT_CHARS:
                    CONSOLE.print(
                        f"[yellow]⚠ Текст завдання скорочено до "
                        f"{MAX_FILE_PROMPT_CHARS} символів, щоб додати файл.[/]"
                    )
                    LOGGER.info("User task was truncated for file attachment")
                elif not pending_attachment:
                    CONSOLE.print(
                        f"[yellow]⚠ Запит скорочено до {MAX_MESSAGE_CHARS} символів "
                        "для безпечного надсилання.[/]"
                    )
                    LOGGER.info("User prompt was truncated to context limit")

            messages.append({"role": "user", "content": prompt_for_context})
            with CONSOLE.status("[bold yellow]Groq думає...[/]", spinner="dots"):
                reply, web_search_used, sources, rate_limits = request_response(
                    client,
                    messages,
                )
            messages.append(
                {"role": "assistant", "content": truncate_text(reply)}
            )
            pending_attachment = None
            try:
                save_chat_history(history_path, messages)
            except OSError as error:
                print_error(f"Не вдалося зберегти історію: {error}")
                LOGGER.exception("Unable to save chat history")
            print_assistant_reply(reply)

            if web_search_used:
                CONSOLE.print(Text("[Використано веб-пошук]", style="bold cyan"))

            if sources:
                print_sources(sources)

            print_rate_limits(rate_limits)

            LOGGER.info(
                "Response received; web_search_used=%s; sources=%s; "
                "rate_limits_received=%s",
                web_search_used,
                len(sources),
                bool(rate_limits),
            )

        except RateLimitError as error:
            messages.pop()
            print_error(f"Перевищено ліміт Groq: {error}", "Ліміт Groq")
            LOGGER.warning("Groq rate limit: %s", error)

        except APIConnectionError:
            messages.pop()
            print_error(
                "Немає з’єднання з Groq. Перевірте інтернет.",
                "З’єднання відсутнє",
            )
            LOGGER.warning("Groq connection error")

        except APIStatusError as error:
            messages.pop()
            print_error(
                f"Помилка Groq ({error.status_code}): {error.message}",
                "Помилка Groq",
            )
            LOGGER.warning("Groq API status error; status=%s", error.status_code)

        except Exception as error:
            messages.pop()
            print_error(f"Неочікувана помилка: {error}", "Неочікувана помилка")
            LOGGER.exception("Unexpected application error")


if __name__ == "__main__":
    run_chat()
