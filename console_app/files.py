"""Safe local text extraction and attachment preparation."""

from __future__ import annotations

from pathlib import Path

from .chat import truncate_text
from .constants import (
    MAX_FILE_BYTES,
    MAX_FILE_PROMPT_CHARS,
    MAX_FILE_TEXT_CHARS,
    MAX_SPREADSHEET_ROWS,
    SUPPORTED_FILE_TYPES,
)


def decode_text_file(file_path: Path) -> str:
    """Read common text files while handling typical Ukrainian encodings."""
    raw_content = file_path.read_bytes()

    if b"\x00" in raw_content:
        raise ValueError("Файл схожий на двійковий, а не текстовий.")

    for encoding in ("utf-8-sig", "utf-8", "cp1251", "utf-16"):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("Не вдалося визначити кодування текстового файла.")


def extract_pdf_text(file_path: Path) -> str:
    """Extract a PDF text layer without uploading the original file."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]

    return "\n\n".join(pages).strip()


def extract_docx_text(file_path: Path) -> str:
    """Extract paragraphs and table cells from a Word document."""
    from docx import Document

    document = Document(file_path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cell for cell in cells if cell))

    return "\n".join(part for part in parts if part).strip()


def extract_xlsx_text(file_path: Path) -> str:
    """Extract cell values from a bounded number of non-empty rows."""
    from openpyxl import load_workbook

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    parts: list[str] = []
    rows_read = 0

    try:
        for worksheet in workbook.worksheets:
            if rows_read >= MAX_SPREADSHEET_ROWS:
                break

            parts.append(f"Аркуш: {worksheet.title}")

            for row in worksheet.iter_rows(values_only=True):
                if rows_read >= MAX_SPREADSHEET_ROWS:
                    break

                values = [str(value).strip() if value is not None else "" for value in row]

                if any(values):
                    parts.append(" | ".join(values))
                    rows_read += 1
    finally:
        workbook.close()

    return "\n".join(parts).strip()


def extract_file_text(file_path: Path) -> str:
    """Choose an extractor only for a supported local file extension."""
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)
    if extension == ".docx":
        return extract_docx_text(file_path)
    if extension == ".xlsx":
        return extract_xlsx_text(file_path)

    return decode_text_file(file_path)


def prepare_file_attachment(file_path: Path) -> str:
    """Validate a local file and return bounded extracted text for the model."""
    if not file_path.exists() or not file_path.is_file():
        raise ValueError("Файл не знайдено або це не звичайний файл.")

    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_FILE_TYPES:
        supported = ", ".join(SUPPORTED_FILE_TYPES)
        raise ValueError(f"Непідтримуваний формат. Доступні: {supported}.")

    file_size = file_path.stat().st_size
    if file_size > MAX_FILE_BYTES:
        raise ValueError(
            f"Файл має {file_size / 1024 / 1024:.1f} МБ. Максимум — "
            f"{MAX_FILE_BYTES // 1024 // 1024} МБ."
        )

    extracted_text = extract_file_text(file_path)
    extracted_text = "\n".join(line.rstrip() for line in extracted_text.splitlines())

    if not extracted_text.strip():
        raise ValueError(
            "Не вдалося отримати текст. Для PDF це може бути скан без "
            "текстового шару."
        )

    attachment_text = truncate_text(extracted_text.strip(), MAX_FILE_TEXT_CHARS)
    return (
        f"[Прикріплений файл: {file_path.name}; формат: {extension}; "
        f"розмір: {file_size / 1024:.1f} КБ]\n"
        "Вміст файла:\n"
        f"{attachment_text}"
    )


def build_file_request(prompt: str, attachment: str) -> str:
    """Keep a user task and a bounded attachment within one API message."""
    task = truncate_text(prompt, MAX_FILE_PROMPT_CHARS)
    return (
        f"Завдання користувача щодо прикріпленого файла:\n{task}\n\n"
        f"{attachment}\n\n"
        "Відповідай за змістом файла. Якщо даних недостатньо, прямо скажи це."
    )
